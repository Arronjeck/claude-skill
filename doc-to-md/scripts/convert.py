#!/usr/bin/env python3
"""将 .docx 或 .pdf 文件转换为可读的 Markdown 格式。"""
import sys
import os
import re
import subprocess
from pathlib import Path


# ─── 依赖安装 ────────────────────────────────────────────────────────────────

def ensure_deps(*packages):
    for pkg in packages:
        import_name = pkg.split("==")[0].replace("-", "_")
        try:
            __import__(import_name)
        except ImportError:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", pkg,
                 "-i", "https://pypi.org/simple/", "-q"],
                stdout=subprocess.DEVNULL,
            )


# ─── DOCX 转换 ───────────────────────────────────────────────────────────────

_NS_A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
_NS_R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_NS_WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_NS_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"


def _extract_para_images(para_elem, doc, assets_dir: Path, img_counter: list) -> list[str]:
    """提取段落中的所有图片，保存到 assets_dir，返回 Markdown 图片引用列表。"""
    refs = []
    drawings = para_elem.findall(f".//{{{_NS_WP}}}inline")
    for drawing in drawings:
        blips = drawing.findall(f".//{{{_NS_A}}}blip")
        for blip in blips:
            r_id = blip.get(f"{{{_NS_R}}}embed")
            if not r_id:
                continue
            try:
                part = doc.part.related_parts[r_id]
            except KeyError:
                continue
            ext = Path(part.partname).suffix or ".png"
            img_counter[0] += 1
            base = f"image{img_counter[0]}"
            filename = f"{base}{ext}"
            dest = assets_dir / filename
            n = 1
            while dest.exists():
                filename = f"{base}({n}){ext}"
                dest = assets_dir / filename
                n += 1
            dest.write_bytes(part.blob)
            # 使用相对路径：仅用 assets_dir 的最后一级目录名
            refs.append(f"![{filename}]({assets_dir.name}/{filename})")
    return refs


def convert_docx(input_path: str, assets_dir: Path | None = None) -> str:
    ensure_deps("python-docx")
    from docx import Document

    # 用 open() 传文件对象，避免 python-docx 在 Windows 上处理中文路径时编码出错
    with open(input_path, "rb") as f:
        doc = Document(f)

    # assets 目录：与输出 md 同级，名为 <文档名>_assets
    if assets_dir is None:
        assets_dir = Path(input_path).with_suffix("") / "_assets"
        assets_dir = Path(str(Path(input_path).with_suffix("")) + "_assets")
    assets_dir.mkdir(parents=True, exist_ok=True)
    img_counter = [0]  # 用列表以便嵌套函数修改

    # 收集所有 run 的字体大小，用于推断标题层级
    all_sizes = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                all_sizes.append(run.font.size)

    if all_sizes:
        from collections import Counter
        body_size = Counter(all_sizes).most_common(1)[0][0]
    else:
        body_size = None

    def para_heading_level(para):
        style = para.style.name if para.style else ""
        text = para.text.strip()

        m = re.match(r"Heading (\d)", style)
        if m:
            return int(m.group(1))

        # 中文编号标题：一、二、三 或 文档标题关键词
        if re.match(r"^[一二三四五六七八九十]+[、.]", text) or re.match(r"^接口文档$", text):
            return 1

        if re.match(r"^\d+\.\s+\S", text) and not re.match(r"^\d+\.\d+", text):
            return 2
        if re.match(r"^\d+\.\d+\S", text):
            return 3

        runs = [r for r in para.runs if r.text.strip()]
        if not runs:
            return 0
        all_bold = all(r.bold for r in runs)
        max_size = max((r.font.size for r in runs if r.font.size), default=None)

        if all_bold and max_size and body_size and max_size > body_size * 1.2:
            return 1
        if all_bold and max_size and body_size and max_size > body_size * 1.05:
            return 2
        if all_bold:
            return 3

        return 0

    def is_code_line(text):
        patterns = [
            r"^(HKEY_|AppStore\.exe|KeyPath:|KeyName:|KeyValue:)",
            r"^[A-Z_]{3,}=",
        ]
        return any(re.match(p, text) for p in patterns)

    body = doc.element.body
    para_idx = 0
    table_idx = 0
    md_lines = []

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            if para_idx >= len(doc.paragraphs):
                para_idx += 1
                continue
            para = doc.paragraphs[para_idx]
            para_idx += 1

            # 先提取图片（图片段落可能同时有文字说明）
            img_refs = _extract_para_images(child, doc, assets_dir, img_counter)
            for ref in img_refs:
                md_lines.append(ref)

            text = para.text.strip()
            if not text:
                if not img_refs:
                    md_lines.append("")
                continue

            level = para_heading_level(para)
            if level == 1:
                md_lines.append(f"\n# {text}")
            elif level == 2:
                md_lines.append(f"\n## {text}")
            elif level == 3:
                md_lines.append(f"\n### {text}")
            elif level == 4:
                md_lines.append(f"\n#### {text}")
            elif is_code_line(text):
                md_lines.append(f"`{text}`")
            else:
                # 保留行内加粗/斜体
                line = ""
                for run in para.runs:
                    t = run.text
                    if not t:
                        continue
                    if run.bold:
                        line += f"**{t}**"
                    elif run.italic:
                        line += f"*{t}*"
                    else:
                        line += t
                md_lines.append(line if line else text)

        elif tag == "tbl":
            if table_idx >= len(doc.tables):
                table_idx += 1
                continue
            table = doc.tables[table_idx]
            table_idx += 1

            rows = []
            for row in table.rows:
                cells = []
                seen = set()
                for cell in row.cells:
                    key = id(cell._tc)
                    if key in seen:
                        continue
                    seen.add(key)
                    cell_text = cell.text.strip().replace("\n", "<br>").replace("|", "\\|")
                    cells.append(cell_text)
                rows.append(cells)

            if not rows:
                continue

            # 单列单行 → 代码块
            if len(rows) == 1 and len(rows[0]) == 1:
                content = rows[0][0].replace("<br>", "\n")
                md_lines.append(f"\n```\n{content}\n```\n")
            else:
                md_lines.append("")
                header = rows[0]
                md_lines.append("| " + " | ".join(header) + " |")
                md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                for row in rows[1:]:
                    # 补齐列数
                    while len(row) < len(header):
                        row.append("")
                    md_lines.append("| " + " | ".join(row[:len(header)]) + " |")
                md_lines.append("")

    result = "\n".join(md_lines)

    # 如果没有提取到任何图片，删除空的 assets 目录
    if img_counter[0] == 0:
        try:
            assets_dir.rmdir()
        except OSError:
            pass
    else:
        print(f"已提取 {img_counter[0]} 张图片到 {assets_dir}/", file=sys.stderr)

    return result


# ─── PDF 转换 ────────────────────────────────────────────────────────────────

def convert_pdf(input_path: str) -> str:
    """
    优先用 pymupdf（fitz）提取带位置信息的文本，推断标题层级。
    fallback 到 pypdf 纯文本提取。
    """
    try:
        ensure_deps("pymupdf")
        import fitz  # noqa: F401
        return _pdf_via_pymupdf(input_path)
    except Exception as e:
        print(f"[警告] pymupdf 不可用（{e}），降级使用 pypdf，表格和代码块识别质量会下降", file=sys.stderr)

    try:
        ensure_deps("pypdf")
        return _pdf_via_pypdf(input_path)
    except Exception as e:
        raise RuntimeError(f"无法解析 PDF：{e}") from e


def _pdf_via_pymupdf(input_path: str) -> str:
    import fitz
    from collections import Counter

    doc = fitz.open(input_path)

    # 收集全文字号分布，推断正文基准字号
    all_sizes = []
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    all_sizes.append(round(span["size"], 1))
    body_size = Counter(all_sizes).most_common(1)[0][0] if all_sizes else 12.0

    # 推断表格左边距基准（最常见的 block x0）
    all_x0 = []
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            if block["type"] == 0:
                all_x0.append(round(block["bbox"][0], 0))
    table_x0 = Counter(all_x0).most_common(1)[0][0] if all_x0 else 0

    def block_fonts(block):
        return {span.get("font", "")
                for line in block["lines"] for span in line["spans"]}

    def block_size(block):
        return {round(span["size"], 1)
                for line in block["lines"] for span in line["spans"]}

    def block_lines_text(block):
        return ["".join(s["text"] for s in line["spans"]).strip()
                for line in block["lines"]]

    def is_code_block(block):
        return any("SourceCodePro" in f or "Courier" in f or "Mono" in f
                   for f in block_fonts(block))

    def is_table_block(block):
        return block_size(block) == {body_size}

    def is_continuation_block(block):
        """单行、x0 明显偏右（续行溢出），应合并到上一行最后一列。"""
        lines = [l for l in block_lines_text(block) if l]
        if len(lines) != 1:
            return False
        x0 = round(block["bbox"][0], 0)
        return x0 > table_x0 + 50  # 偏右超过 50pt 视为续行

    def is_heading(text, size):
        if re.match(r"^\d+\.\d+", text) or re.match(r"^\d+\.\s+\S", text):
            return True
        return (size / body_size if body_size else 1) >= 1.8

    # 将所有页的 text block 拍平为一个列表（跨页连续处理）
    all_blocks = []
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            if block["type"] == 0:
                all_blocks.append(block)

    md_lines = []
    i = 0
    while i < len(all_blocks):
        block = all_blocks[i]

        # ── 代码块：合并连续 SourceCodePro block ────────────────
        if is_code_block(block):
            code_lines = []
            while i < len(all_blocks) and is_code_block(all_blocks[i]):
                for lt in block_lines_text(all_blocks[i]):
                    if lt:
                        code_lines.append(lt)
                i += 1
            md_lines.append(f"\n```\n" + "\n".join(code_lines) + "\n```\n")
            continue

        # ── 表格：合并连续 table block（含跨页、续行）───────────
        if is_table_block(block):
            table_rows = []
            while i < len(all_blocks) and is_table_block(all_blocks[i]):
                b = all_blocks[i]
                cells = [lt for lt in block_lines_text(b) if lt]
                if not cells:
                    i += 1
                    continue
                # 续行：x0 偏右且单行 → 追加到上一行最后一列
                if table_rows and is_continuation_block(b):
                    table_rows[-1][-1] += cells[0]
                else:
                    table_rows.append(cells)
                i += 1

            if not table_rows:
                continue

            col_count = max(len(r) for r in table_rows)
            if col_count >= 2:
                rows_padded = [r + [""] * (col_count - len(r)) for r in table_rows]
                md_lines.append("")
                md_lines.append("| " + " | ".join(rows_padded[0]) + " |")
                md_lines.append("| " + " | ".join(["---"] * col_count) + " |")
                for row in rows_padded[1:]:
                    md_lines.append("| " + " | ".join(row) + " |")
                md_lines.append("")
            else:
                for row in table_rows:
                    md_lines.append(row[0] if row else "")
            continue

        # ── 普通文本 / 标题 ──────────────────────────────────────
        for line in block["lines"]:
            line_text = "".join(s["text"] for s in line["spans"]).strip()
            if not line_text:
                continue
            max_size = max(round(s["size"], 1) for s in line["spans"])
            if is_heading(line_text, max_size):
                if re.match(r"^\d+\.\d+", line_text):
                    md_lines.append(f"\n### {line_text}")
                elif re.match(r"^\d+\.\s+\S", line_text):
                    md_lines.append(f"\n## {line_text}")
                else:
                    md_lines.append(f"\n# {line_text}")
            else:
                md_lines.append(line_text)
        i += 1

    return "\n".join(md_lines)


def _pdf_via_pypdf(input_path: str) -> str:
    import pypdf

    reader = pypdf.PdfReader(input_path)
    md_lines = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        lines = text.splitlines()
        for line in lines:
            line = line.strip()
            if not line:
                md_lines.append("")
                continue
            # 数字编号章节标题：如 "1. xxx"、"2.3xxx"
            if re.match(r"^\d+\.\s+\S", line) and not re.match(r"^\d+\.\d+", line):
                md_lines.append(f"\n## {line}")
            elif re.match(r"^\d+\.\d+\S", line):
                md_lines.append(f"\n### {line}")
            # 文档标题：首页第一个非空短行（不含冒号/数字）
            elif i == 0 and not md_lines and len(line) < 40 and not re.search(r"[：:。，,]", line):
                md_lines.append(f"\n# {line}")
            else:
                md_lines.append(line)
        md_lines.append("")

    return "\n".join(md_lines)


# ─── 入口 ────────────────────────────────────────────────────────────────────

def _resolve_path(input_path: str) -> tuple[str, str]:
    """
    On Windows with GBK filesystem + UTF-8 Python, CLI-passed Chinese paths
    may not match os.scandir filenames. Returns (real_path, display_path):
      real_path    - path that open() can use (may be garbled string)
      display_path - path for deriving output filename (original input)
    """
    if os.path.exists(input_path):
        return input_path, input_path

    parent = os.path.dirname(input_path) or "."
    target = os.path.basename(input_path)
    suffix = os.path.splitext(target)[1].lower()

    try:
        candidates = [e for e in os.scandir(parent)
                      if e.name.lower().endswith(suffix)]
    except OSError:
        return input_path, input_path

    if len(candidates) == 1:
        return candidates[0].path, input_path

    def ascii_overlap(a, b):
        sa = set(c for c in a if ord(c) < 128)
        sb = set(c for c in b if ord(c) < 128)
        return len(sa & sb)

    best = max(candidates, key=lambda e: ascii_overlap(e.name, target))
    return best.path, input_path


def convert(input_path: str, output_path: str | None = None) -> str:
    real_path, display_path = _resolve_path(input_path)
    p_real = Path(real_path)
    p_display = Path(display_path)
    suffix = p_real.suffix.lower()

    # 输出路径用 display_path 派生，保留原始（UTF-8）文件名
    if output_path is None:
        output_path = str(p_display.with_suffix(".md"))

    out_p = Path(output_path)
    assets_dir = out_p.parent / (out_p.stem + "_assets")

    if suffix == ".docx":
        content = convert_docx(real_path, assets_dir)
    elif suffix == ".pdf":
        content = convert_pdf(real_path)
    else:
        raise ValueError(f"不支持的文件类型：{suffix}（仅支持 .docx 和 .pdf）")

    # 压缩多余空行（超过 2 个连续空行合并为 1 个）
    content = re.sub(r"\n{3,}", "\n\n", content)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)

    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python convert.py <input.docx|input.pdf> [output.md]")
        sys.exit(1)
    out = convert(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
    print(f"已生成: {out}")
